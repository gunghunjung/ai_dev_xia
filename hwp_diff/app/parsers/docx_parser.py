"""Parser for DOCX format using python-docx."""
import uuid
from pathlib import Path
from typing import List, Optional

from app.models.document import (
    DocumentStructure, DocumentBlock, TextSpan, TableData, TableCell
)
from app.parsers.base_parser import BaseParser
from app.utils.logger import get_logger
from app.utils.text_utils import detect_numbering

logger = get_logger("parsers.docx")


class DOCXParser(BaseParser):
    """Parses Microsoft Word DOCX files using python-docx."""

    def parse(self, file_path: str) -> DocumentStructure:
        p = self._check_file(file_path)
        doc_id = self._new_doc_id(file_path)
        logger.info("Parsing DOCX: %s", file_path)

        doc = DocumentStructure(
            doc_id=doc_id,
            file_path=str(p),
            file_type="docx",
            parse_confidence=0.95,
            page_confidence=0.3,
            table_confidence=0.95,
        )

        try:
            import docx
            from docx.oxml.ns import qn

            word_doc = docx.Document(str(p))

            para_idx = 0
            table_idx = 0

            # Iterate body elements in document order (paragraphs and tables)
            # python-docx exposes word_doc.paragraphs and word_doc.tables separately,
            # but we need document order. We iterate the body XML.
            body = word_doc.element.body
            for child in body:
                tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

                if tag == "p":
                    # Wrap in docx Paragraph
                    try:
                        from docx.text.paragraph import Paragraph
                        para = Paragraph(child, word_doc)
                    except Exception:
                        continue

                    block = self._parse_paragraph(para, para_idx)
                    if block is not None:
                        doc.blocks.append(block)
                        para_idx += 1

                elif tag == "tbl":
                    try:
                        from docx.table import Table
                        table = Table(child, word_doc)
                    except Exception:
                        continue

                    block = self._parse_table(table, table_idx, para_idx)
                    if block is not None:
                        doc.blocks.append(block)
                        table_idx += 1
                        para_idx += 1

            doc.total_pages = max(1, para_idx // 40 + 1)
            doc.title = self._extract_title(doc.blocks)
            logger.info(
                "DOCX parsed: %d blocks, %d pages est.",
                len(doc.blocks), doc.total_pages
            )

        except ImportError:
            logger.error("python-docx is not installed. Cannot parse DOCX.")
            doc.parse_confidence = 0.0
        except Exception as e:
            logger.exception("Error parsing DOCX %s: %s", file_path, e)
            doc.parse_confidence = 0.1

        return doc

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _heading_level_from_style(self, style_name: str) -> int:
        """Determine heading level from paragraph style name."""
        if not style_name:
            return 0
        sn = style_name.lower().strip()
        # English headings: "Heading 1", "Heading1"
        import re
        m = re.search(r"heading\s*(\d)", sn)
        if m:
            return int(m.group(1))
        # Korean headings
        m = re.search(r"제목\s*(\d)", sn)
        if m:
            return int(m.group(1))
        if sn in ("heading", "제목", "title", "제목1"):
            return 1
        if "제목" in sn or "heading" in sn:
            return 1
        return 0

    def _parse_paragraph(self, para, para_idx: int) -> Optional[DocumentBlock]:
        """Parse a docx Paragraph into a DocumentBlock."""
        content = para.text.strip()
        if not content:
            return None

        style_name = ""
        try:
            style_name = para.style.name if para.style else ""
        except Exception:
            pass

        heading_level = self._heading_level_from_style(style_name)

        # Build spans from runs
        spans: List[TextSpan] = []
        for run in para.runs:
            if run.text:
                span = TextSpan(
                    text=run.text,
                    bold=bool(run.bold),
                    italic=bool(run.italic),
                    underline=bool(run.underline),
                    font_size=float(run.font.size.pt) if run.font.size else 10.0,
                    font_name=run.font.name or "",
                )
                spans.append(span)

        numbering = detect_numbering(content)

        if heading_level > 0:
            block_type = "heading"
        elif numbering:
            block_type = "list_item"
        else:
            block_type = "paragraph"

        block_id = f"p_{para_idx}_{uuid.uuid4().hex[:6]}"
        page_hint = para_idx // 40 + 1

        return DocumentBlock(
            block_id=block_id,
            block_type=block_type,
            content=content,
            spans=spans,
            level=heading_level,
            paragraph_index=para_idx,
            page_hint=page_hint,
            numbering=numbering,
            style_name=style_name,
        )

    def _parse_table(self, table, table_idx: int, para_idx: int) -> Optional[DocumentBlock]:
        """Parse a docx Table into a DocumentBlock with TableData."""
        table_id = f"tbl_{table_idx}"
        rows_data: List[List[TableCell]] = []
        max_cols = 0

        for r_idx, row in enumerate(table.rows):
            row_cells: List[TableCell] = []
            for c_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                # Attempt to get span info
                try:
                    from docx.oxml.ns import qn
                    tc = cell._tc
                    tcPr = tc.find(qn("w:tcPr"))
                    row_span = 1
                    col_span = 1
                    if tcPr is not None:
                        gridSpan = tcPr.find(qn("w:gridSpan"))
                        if gridSpan is not None:
                            col_span = int(gridSpan.get(qn("w:val"), 1))
                except Exception:
                    row_span = 1
                    col_span = 1

                tc_obj = TableCell(
                    row=r_idx,
                    col=c_idx,
                    row_span=row_span,
                    col_span=col_span,
                    content=cell_text,
                )
                row_cells.append(tc_obj)
            if row_cells:
                rows_data.append(row_cells)
                max_cols = max(max_cols, len(row_cells))

        if not rows_data:
            return None

        table_data = TableData(rows=len(rows_data), cols=max_cols, cells=rows_data)
        flat_text = "\n".join(
            " | ".join(cell.content for cell in row) for row in rows_data
        )
        block_id = f"tbl_{table_idx}_{uuid.uuid4().hex[:6]}"
        page_hint = para_idx // 40 + 1

        return DocumentBlock(
            block_id=block_id,
            block_type="table",
            content=flat_text,
            paragraph_index=para_idx,
            page_hint=page_hint,
            table_data=table_data,
            table_id=table_id,
        )

    def _extract_title(self, blocks: List[DocumentBlock]) -> str:
        """Extract document title from heading blocks."""
        for block in blocks:
            if block.block_type == "heading" and block.level == 1:
                return block.content
        for block in blocks:
            if block.content.strip():
                return block.content[:80]
        return ""
